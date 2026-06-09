"""PDF, DOCX, and Excel document readers for the Nexus tool library (H44)."""

from __future__ import annotations

import asyncio
from datetime import datetime
from pathlib import Path
from typing import Any

from nexus.core.errors import ToolError
from nexus.tools.library.document_chunker import DocumentChunker
from nexus.tools.library.document_types import DocumentChunk, DocumentMetadata, ParsedDocument

# ---------------------------------------------------------------------------
# PDF readers
# ---------------------------------------------------------------------------


def _pdf_metadata(path: Path, info: dict[str, Any], parser: str) -> DocumentMetadata:
    return DocumentMetadata(
        source=str(path),
        format="pdf",
        parser=parser,
        title=info.get("title") or None,
        author=info.get("author") or None,
    )


def _read_pdf_fitz(path: Path, chunker: DocumentChunker) -> ParsedDocument:
    import fitz

    doc = fitz.open(str(path))
    info: dict[str, Any] = doc.metadata or {}
    meta = _pdf_metadata(path, info, "pymupdf")
    meta.total_pages = len(doc)
    all_chunks: list[DocumentChunk] = []
    raw_parts: list[str] = []
    for page_num, page in enumerate(doc):
        text: str = page.get_text("text")
        if text.strip():
            raw_parts.append(text)
            all_chunks.extend(chunker.chunk(text, meta, page=page_num + 1))
    doc.close()
    return ParsedDocument(chunks=all_chunks, metadata=meta, raw_text="\n\n".join(raw_parts))


def _read_pdf_pypdf(path: Path, chunker: DocumentChunker) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pdf_info = reader.metadata
    info: dict[str, Any] = {}
    if pdf_info:
        info = {
            "title": getattr(pdf_info, "title", None),
            "author": getattr(pdf_info, "author", None),
        }
    meta = _pdf_metadata(path, info, "pypdf")
    meta.total_pages = len(reader.pages)
    all_chunks: list[DocumentChunk] = []
    raw_parts: list[str] = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            raw_parts.append(text)
            all_chunks.extend(chunker.chunk(text, meta, page=page_num + 1))
    return ParsedDocument(chunks=all_chunks, metadata=meta, raw_text="\n\n".join(raw_parts))


async def read_pdf(path: str | Path, chunker: DocumentChunker) -> ParsedDocument:
    """Read a PDF file using PyMuPDF (preferred) or pypdf (fallback).

    Raises:
        ToolError: With code MISSING_DEPENDENCY when neither library is available.
    """
    use_fitz = False
    try:
        import fitz  # noqa: F401  # type: ignore[import]

        use_fitz = True
    except ImportError:
        try:
            from pypdf import PdfReader  # noqa: F401  # type: ignore[import]
        except ImportError:
            raise ToolError(
                "PDF support requires: pip install 'nexus-ai[documents]'",
                code="MISSING_DEPENDENCY",
            ) from None
    reader_fn = _read_pdf_fitz if use_fitz else _read_pdf_pypdf
    return await asyncio.to_thread(reader_fn, Path(path), chunker)


# ---------------------------------------------------------------------------
# DOCX reader
# ---------------------------------------------------------------------------


def _heading_level(style_name: str | None) -> int:
    """Return heading depth for a Word style name, 0 for body text."""
    if not style_name or not style_name.startswith("Heading "):
        return 0
    try:
        return int(style_name[8:])
    except ValueError:
        return 0


def _table_to_markdown(table: Any) -> str:
    """Render a python-docx Table as a Markdown table string."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * len(rows[0])) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, sep] + body)


def _flush_section(
    texts: list[str],
    meta: DocumentMetadata,
    section_path: list[str],
    chunker: DocumentChunker,
    all_chunks: list[DocumentChunk],
    raw_parts: list[str],
) -> None:
    if not texts:
        return
    joined = "\n".join(texts)
    raw_parts.append(joined)
    all_chunks.extend(chunker.chunk(joined, meta, section_path=list(section_path)))
    texts.clear()


def _read_docx_sync(path: Path, chunker: DocumentChunker) -> ParsedDocument:
    import docx

    doc = docx.Document(str(path))
    props = doc.core_properties
    raw_created = props.created
    created: datetime | None = raw_created if isinstance(raw_created, datetime) else None
    meta = DocumentMetadata(
        source=str(path),
        format="docx",
        parser="python-docx",
        title=props.title or None,
        author=props.author or None,
        created_at=created,
    )
    all_chunks: list[DocumentChunk] = []
    raw_parts: list[str] = []
    heading_stack: list[str] = []
    current_path: list[str] = []
    current_texts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        level = _heading_level(style_name)
        if level > 0:
            _flush_section(current_texts, meta, current_path, chunker, all_chunks, raw_parts)
            heading_stack = heading_stack[: level - 1] + [text]
            current_path = list(heading_stack)
        else:
            current_texts.append(text)

    for table in doc.tables:
        md = _table_to_markdown(table)
        if md:
            current_texts.append(md)

    _flush_section(current_texts, meta, current_path, chunker, all_chunks, raw_parts)
    return ParsedDocument(chunks=all_chunks, metadata=meta, raw_text="\n\n".join(raw_parts))


async def read_docx(path: str | Path, chunker: DocumentChunker) -> ParsedDocument:
    """Read a Word (.docx) file using python-docx.

    Raises:
        ToolError: With code MISSING_DEPENDENCY when python-docx is not installed.
    """
    try:
        import docx  # noqa: F401  # type: ignore[import]
    except ImportError:
        raise ToolError(
            "DOCX support requires: pip install 'nexus-ai[documents]'",
            code="MISSING_DEPENDENCY",
        ) from None
    return await asyncio.to_thread(_read_docx_sync, Path(path), chunker)


# ---------------------------------------------------------------------------
# Excel reader
# ---------------------------------------------------------------------------

_EXCEL_MAX_ROWS = 1000


def _sheet_to_markdown(sheet: Any) -> str:
    """Render an openpyxl worksheet as a Markdown table, capped at 1000 rows."""
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return ""
    truncated = len(rows) > _EXCEL_MAX_ROWS
    rows = rows[:_EXCEL_MAX_ROWS]
    header = "| " + " | ".join(str(v or "") for v in rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * len(rows[0])) + " |"
    body = ["| " + " | ".join(str(v or "") for v in row) + " |" for row in rows[1:]]
    text = "\n".join([header, sep] + body)
    if truncated:
        text += "\n\n*(truncated — showing first 1000 rows)*"
    return text


def _read_excel_sync(path: Path, chunker: DocumentChunker) -> ParsedDocument:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    title = (wb.properties.title or None) if wb.properties else None
    meta = DocumentMetadata(source=str(path), format="xlsx", parser="openpyxl", title=title)
    meta.total_sheets = len(wb.sheetnames)
    all_chunks: list[DocumentChunk] = []
    raw_parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        md = _sheet_to_markdown(ws)
        if md.strip():
            raw_parts.append(md)
            all_chunks.extend(chunker.chunk(md, meta, section_path=[sheet_name], sheet=sheet_name))

    wb.close()
    return ParsedDocument(chunks=all_chunks, metadata=meta, raw_text="\n\n".join(raw_parts))


async def read_excel(path: str | Path, chunker: DocumentChunker) -> ParsedDocument:
    """Read an Excel (.xlsx) file using openpyxl.

    Raises:
        ToolError: With code MISSING_DEPENDENCY when openpyxl is not installed.
    """
    try:
        import openpyxl  # noqa: F401  # type: ignore[import]
    except ImportError:
        raise ToolError(
            "Excel support requires: pip install 'nexus-ai[documents]'",
            code="MISSING_DEPENDENCY",
        ) from None
    return await asyncio.to_thread(_read_excel_sync, Path(path), chunker)
